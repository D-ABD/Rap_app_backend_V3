# rap_app/utils/exporter.py

import io
import csv

from django.http import HttpResponse
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4


class Exporter:
    """
    🧰 Classe utilitaire générique pour exporter un queryset Django dans plusieurs formats :
    - CSV
    - Word (.docx)
    - PDF (via ReportLab)

    Fonctionnalité
    ==============
    Permet l'extraction, la transformation et la sérialisation d'un queryset Django en plusieurs formats
    bureautiques standards, en s'appuyant sur des champs explicitement listés et un mapping éventuel
    d'en-têtes. L'utilisateur peut facilement brancher cette classe sur n'importe quel queryset ou
    liste d'objets Django.

    Contrat technique
    ================
    - Construction : Exporter(queryset, fields: list[str], headers: Optional[list[str]])
      - queryset : QuerySet Django (ou potentiellement tout itérable d'objets dotés des attributs attendus)
      - fields : liste de noms d'attributs ou chemins d'attributs imbriqués (ex: "formation.nom")
      - headers : Noms des colonnes à afficher à l'export (par défaut identique à fields)

    - Méthode export_csv : str (nom de fichier) -> HttpResponse contenant un flux CSV
    - Méthode export_word : str (nom de fichier) -> HttpResponse contenant un .docx
    - Méthode export_pdf : str (nom de fichier) -> HttpResponse contenant un PDF

    Usage exemple
    =============
    >>> exporter = Exporter(queryset, fields=["nom", "created_by.username"], headers=["Nom", "Auteur"])
    >>> response = exporter.export_csv("export.csv")
    >>> # response est une HttpResponse à renvoyer telle quelle dans une vue Django

    Limites et dépendances
    =====================
    - Dépend explicitement de la présence sur chaque objet du queryset des attributs listés dans `fields`
    - Résolution dynamique sur les relations Django via le séparateur "." (cf. _resolve_field)
    - Nécessite les bibliothèques externes Document (python-docx) et reportlab pour la sortie Word et PDF
    - Pas de vérification de type ou de format sur les objets; toute absence d'attribut produit une chaîne vide ("")
    - Renvoie une HttpResponse Django prête pour le download, ne gère pas l'appel en dehors de ce contexte

      - Manipulations complexes de types personnalisés (dates, choices, etc.) non automatisées
      - Les performances sur de très gros querysets ne sont pas garanties (pas de pagination/streaming)
    """

    def __init__(self, queryset, fields, headers=None):
        """
        Initialise l'exporteur générique.

        Fonctionnalité :
            Stocke la référence au queryset à exporter et la liste des champs à extraire.
            Gère l'éventuel mapping personnalisé des entêtes.

        Contrat technique :
            queryset : QuerySet Django (ou équivalent itérable d'objets)
            fields : list[str] — Ex : ["nom", "created_by.username"]
            headers : list[str] optionnelle — même longueur que fields, sinon headers = fields

        Exemple
        -------
            Exporter(qs, fields=["code", "auteur.email"], headers=["Code interne", "Mail de l'auteur"])

        Limites/dépendances :
            - La présence des champs/d’attributs listés dans fields est impérative (sinon vide silencieux)
            - Les modèles/objets peuvent être de n'importe quel type compatible
        """
        self.queryset = queryset
        self.fields = fields
        self.headers = headers or fields

    def _resolve_field(self, obj, field_path):
        """
        Résout dynamiquement les relations et attributs pour un objet donné.

        Fonctionnalité :
            Navigue de façon récursive dans les attributs de l’objet,
            pour supporter les chemins du type "auteur.nom" ou "formation.manager.email".
            Si l'attribut rencontré est une méthode, elle est appelée sans argument.

        Contrat technique :
            obj : objet sur lequel on navigue
            field_path : str, chemin d'attributs séparés par "."
            retourne : str (représentation textuelle finale, ou "" si valeur None/non trouvée)

        Exemple
        -------
            Soit obj ayant obj.auteur.nom = "Alice"
            _resolve_field(obj, "auteur.nom")  # renvoie "Alice"

        Limites/dépendances :
            - S'appuie sur getattr(obj, part, None) et appelle la cible si c'est une méthode (sans arguments)
            - Si un niveau est absent : retourne ""
            - Ne gère pas les propriétés nécessitant des arguments ou des propriétés @property complexes
        """
        for part in field_path.split("."):
            obj = getattr(obj, part, None)
            if callable(obj):
                obj = obj()
        return str(obj) if obj is not None else ""

    def get_data(self):
        """
        Génère les données ligne par ligne à partir du queryset pour l'export.

        Fonctionnalité :
            Pour chaque objet du queryset, construit une liste de valeurs correspondant
            aux champs/directions définis par self.fields, prêtes à être écrites dans le format cible.

        Contrat technique :
            retourne : générateur qui yield list[str] (une liste de valeurs, une par champ)

        Exemple
        -------
            self.fields = ["nom", "created_by.email"]
            Pour chaque obj avec obj.nom = "A", obj.created_by.email = "a@x.fr"     
            yield ["A", "a@x.fr"]

        Limites/dépendances :
            - Attendu : les objets du queryset sont dotés des attributs correpondants aux fields listés
            - Les valeurs sont toutes castées en str ("None" => "") pour simplification
        """
        for obj in self.queryset:
            yield [self._resolve_field(obj, f) for f in self.fields]

    def export_csv(self, filename="export.csv"):
        """
        Exporte les données en CSV sous forme de HttpResponse Django.

        Fonctionnalité :
            Génère un fichier CSV avec la première ligne comme en-tête puis une ligne par objet du queryset.

        Contrat technique :
            filename : str, nom du fichier à télécharger (ex: export.csv)
            retourne  : HttpResponse (content_type="text/csv") — pour usage direct dans une réponse Django

        Exemple
        -------
            response = exporter.export_csv("etudiants_mai.csv")
            # response est l'objet HttpResponse, non le contenu du fichier

        Limites/dépendances :
            - Ecriture des données via le module Python standard csv.writer
            - Ne supporte pas explicitement les caractères non-ASCII (encoding dépend de l'instance Python)
            - La longueur des headers doit correspondre à celle des fields
            - Le séparateur est la virgule, non paramétrable ici
        """
        buffer = io.StringIO()
        writer = csv.writer(buffer)
        writer.writerow(self.headers)
        for row in self.get_data():
            writer.writerow(row)
        buffer.seek(0)
        return HttpResponse(
            buffer.getvalue(),
            content_type="text/csv",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    def export_word(self, filename="export.docx"):
        """
        Exporte les données en fichier Word (.docx) téléchargeable.

        Fonctionnalité :
            Génère un document Word avec un titre par défaut ("Export de données"), insère une table
            dont la première ligne est l’en-tête puis autant de lignes que l’export contient d’entrées.

        Contrat technique :
            filename : str (par défaut export.docx)
            retourne : HttpResponse (content_type défini pour Word)

        Exemple
        -------
            response = exporter.export_word("rapport.docx")
            # response à transmettre dans une vue Django

        Limites/dépendances :
            - Utilise la bibliothèque python-docx (Document)
            - N’effectue aucun traitement de style avancé, la table est basique
            - Aucune mise en page ni typographie personnalisée
            - Les valeurs sont toutes converties en str, potentielle perte d'information ou formatage
        """
        doc = Document()
        doc.add_heading("Export de données", 0)

        table = doc.add_table(rows=1, cols=len(self.headers))
        for i, h in enumerate(self.headers):
            table.cell(0, i).text = h

        for row in self.get_data():
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = val

        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        return HttpResponse(
            f.read(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )

    def export_pdf(self, filename="export.pdf"):
        """
        Exporte les données en PDF téléchargeable.

        Fonctionnalité :
            Génère un document PDF (via ReportLab) commençant par un titre, puis inscrit chaque champ
            ligne par ligne. Gère un découpage simple si le texte est trop long. Passe à la page suivante
            si besoin.

        Contrat technique :
            filename : str (par défaut export.pdf)
            retourne : HttpResponse avec application/pdf (pour téléchargement direct)

        Exemple
        -------
            response = exporter.export_pdf("donnees.pdf")
            # response à transmettre dans une vue Django

        Limites/dépendances :
            - Utilise reportlab pour la génération PDF : dépendances externes nécessaires
            - Mise en page minimaliste (texte brut, pas de colonnes ni de tableaux)
            - Les textes longs sont découpés en tranches de 90 caractères, sans gestion de retour à la ligne intelligente
            - Ne gère pas la pagination de façon avancée (simple gestion de l'espace vertical)
            - Aucune gestion de l'encodage/détection de caractères spéciaux
            - headers doivent être cohérents avec fields

            - Améliorer le support multilingue/unicode si besoin
            - Permettre des gabarits de rendu PDF plus avancés
        """
        buffer = io.BytesIO()
        p = canvas.Canvas(buffer, pagesize=A4)
        width, height = A4
        margin = 40
        y = height - margin

        def draw_header():
            """
            Affiche l'en-tête du PDF ("Export de données") en haut de page.

            Fonctionnalité :
                Inscrit un titre en caractères gras à une position fixe sur la page.

            Contrat technique :
                Modifie y (coordonnée verticale), retourne la nouvelle valeur de y
                Sans argument, scope local au rendu du PDF

            Limites/dépendances :
                - Doit être appelé à chaque nouvelle page manuellement
                - La chaîne du titre est en dur ("Export de données")
            """
            p.setFont("Helvetica-Bold", 14)
            p.drawString(margin, y, "Export de données")
            return y - 30

        def draw_row(row_data):
            """
            Inscrit les rubriques d'une ligne d'objet sous forme 'header : valeur'

            Fonctionnalité :
                Parcourt la liste row_data et pour chaque champ, découpe les valeurs longues et
                inscrit chaque sous-partie dans le PDF, déclenchant un saut de page si nécessaire.

            Contrat technique :
                row_data : list[str]
                Aucun retour explicite, effet de bord sur le canvas PDF et la variable y

            Exemple
            -------
                row_data = ['Alice', 'alice@email.fr']
                # Produit "Nom : Alice" puis "Email : alice@email.fr" ligne à ligne

            Limites/dépendances :
                - Utilise la fonction split_text définie localement pour la découpe longue
                - Tout débordement vertical déclenche un saut de page « basique »
                - Pas de gestion des tableaux, chaque champ apparaît en mode "champ: valeur"
            """
            nonlocal y
            p.setFont("Helvetica", 10)
            for i, val in enumerate(row_data):
                # Découpage du texte long
                lines = split_text(val, max_length=90)
                for line in lines:
                    if y < 50:
                        p.showPage()
                        y = height - margin
                    p.drawString(margin, y, f"{self.headers[i]} : {line}")
                    y -= 15
            y -= 10  # espacement entre objets

        def split_text(text, max_length=90):
            """
            Découpe une chaîne longue en morceaux de maximum max_length caractères.

            Fonctionnalité :
                Pour éviter les débordements, les champs trop longs sont découpés
                en "tranches" de taille fixée (par défaut 90 caractères).

            Contrat technique :
                text : str (entrée)
                max_length : int (par défaut 90)
                retourne : list[str] (chaînes découpées)

            Exemple
            -------
                split_text("abcdefghi", 3)  # ['abc', 'def', 'ghi']

            Limites/dépendances :
                - Pas de gestion multilingue ni découpe sur mots
                - Coupe la chaîne sans tenir compte des espaces ni du contenu
            """
            return [text[i:i+max_length] for i in range(0, len(text), max_length)]

        y = draw_header()
        for row in self.get_data():
            if y < 100:
                p.showPage()
                y = height - margin
                y = draw_header()
            draw_row(row)

        p.save()
        buffer.seek(0)
        return HttpResponse(
            buffer.read(),
            content_type="application/pdf",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Content-Length": str(buffer.getbuffer().nbytes),
            }
        )
